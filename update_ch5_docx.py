# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

path = r'E:\develop\BSDP-Bike Sharing Demand Prediction Based on LSTM Model\文本记录内容\终期\x论文.docx'
out_path = r'E:\develop\BSDP-Bike Sharing Demand Prediction Based on LSTM Model\文本记录内容\终期\x论文-第5章已调整.docx'
doc = Document(path)


def insert_paragraph_after(paragraph, text='', style=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para

updates = {
    605: '5.2 系统核心页面与功能实现',
    606: '下面结合平台中的主要业务页面，对系统核心功能的实际实现情况进行说明。考虑到本课题所实现的平台不仅包含后端业务逻辑，还强调前端页面对预测结果、调度信息和运维状态的综合展示，因此第5章在展开时按照“总览感知—系统支撑—预测分析—调度辅助—运维管理”的顺序进行组织。这样既能够突出系统首页在全局监控中的核心地位，也能够更清晰地体现各功能模块之间由数据到预测、由预测到调度、再由调度到运维反馈的业务衔接关系。',
    607: '5.2.2 用户权限与系统管理实现',
    623: '5.2.3 数据处理与需求预测功能实现',
    638: '5.2.4 调度监控与任务管理实现',
    649: '5.2.5 运维管理功能实现',
    672: '本章围绕校园共享单车调度预测与运维管理平台的工程实现情况展开论述，从开发与运行环境、总览页面、系统支撑功能、需求预测页面、调度监控页面以及运维管理页面等多个层面，对系统的实际完成情况进行了说明。可以看出，平台已经完成了系统总览、用户权限管理、参数配置、站点供需状态预测、调度建议生成、站点与车辆运维管理以及日志和备份等主要功能，并在此基础上形成了较为完整的图表、地图和状态联动展示体系。',
    673: '从实现效果看，系统不仅完成了后端业务逻辑的组织，也通过页面化方式将全局运行态势、站点预测结果、调度建议和运维状态较直观地呈现出来，使平台兼具展示性与可操作性。总体而言，第5章所展示的实现内容，已经能够较好支撑论文题目中“调度预测与运维管理平台”的表述，也为后续系统测试与结果分析提供了实际基础。'
}

for idx, text in updates.items():
    doc.paragraphs[idx - 1].text = text

anchor = doc.paragraphs[606 - 1]
style_h = doc.paragraphs[607 - 1].style
style_n = doc.paragraphs[608 - 1].style

p1 = insert_paragraph_after(anchor, '5.2.1 系统总览页面实现', style_h)
p2 = insert_paragraph_after(p1, '系统总览页面是整个平台的首页，也是用户进入系统后感知整体运行状态的核心入口。该页面的设计重点不是单一功能操作，而是通过对关键业务指标、站点状态分布和重点风险对象的集中展示，帮助管理人员和运维人员在较短时间内完成对校园共享单车运行态势的整体判断。因此，在第5章的页面实现中，总览页面应被视为连接预测、调度和运维三个业务模块的重要枢纽。', style_n)
p3 = insert_paragraph_after(p2, '从页面组成上看，总览页面主要由全局 KPI 指标区、校园站点状态地图、重点站点列表以及单站点历史对比展示区构成。其中，KPI 指标区用于汇总站点总数、车辆总数、正常车辆数、重点缺口站点数等全局统计信息；校园站点状态地图用于展示 62 个站点在空间分布上的当前状态；重点站点列表用于突出当前需要优先关注的异常或缺口站点；单站点历史对比展示区则用于补充展示典型站点的历史运行变化情况。通过这种“指标 + 地图 + 列表 + 图表”的组合方式，页面能够兼顾宏观感知与局部分析两类需求。', style_n)
p4 = insert_paragraph_after(p3, '在实现逻辑上，总览页面并不是静态展示，而是依赖后端对站点库存、车辆状态、供需分类结果和统计指标进行统一汇总后，再以图表和地图形式进行前端渲染。页面首次加载时由 Django 完成基础结构输出，随后前端再通过异步方式获取最新业务数据并对 KPI 卡片、地图标记颜色和重点站点信息进行局部刷新。这样既保证了首页展示的直观性，又增强了系统在实际运行过程中的动态监控能力，也为用户继续进入需求预测、调度监控和运维管理页面提供了明确的数据入口。', style_n)

# Apply heading styles again in case plain text reset them
for idx in [605, 607, 623, 638, 649]:
    doc.paragraphs[idx - 1].style = style_h if idx != 605 else doc.paragraphs[605 - 2].style

doc.save(out_path)
print(out_path)
