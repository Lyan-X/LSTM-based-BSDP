# -*- coding: utf-8 -*-
from docx import Document

path = r'E:\develop\BSDP-Bike Sharing Demand Prediction Based on LSTM Model\文本记录内容\终期\x论文-第5章已调整.docx'
doc = Document(path)

normal = doc.paragraphs[595].style
heading = doc.paragraphs[606].style

# Fix mis-styled paragraphs after insertion
for idx in [623, 638, 649]:
    doc.paragraphs[idx - 1].style = normal

# Ensure real subsection headings keep heading style
for idx in [607, 611, 627, 642, 653]:
    doc.paragraphs[idx - 1].style = heading

# Fix duplicated figure numbers in 运维管理部分
replacements = {
    657: '图5-13 站点运维页面整体截图',
    659: '图5-14 站点运维明细表截图',
    662: '图5-15 车辆运维页面整体截图',
    664: '图5-16 车辆运维明细表截图',
    666: '图5-17 单车轨迹历史展开截图',
    668: '图5-18 车辆状态修改操作截图（1）',
    671: '图5-19 车辆状态修改操作截图（2）',
    673: '图5-20 车辆状态修改操作截图（3）',
}
for idx, text in replacements.items():
    doc.paragraphs[idx - 1].text = text

doc.save(path)
print('fixed')
